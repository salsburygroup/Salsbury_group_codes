import os
import argparse
from docx import Document
from docx.shared import Inches
from rdkit import Chem
from rdkit.Chem import Draw, AllChem
import pandas as pd
from io import BytesIO
from matplotlib.backends.backend_pdf import PdfPages
import matplotlib.pyplot as plt
import linecache

def read_molecules(directory, ext, library_file):
    molecules = []
    library_lines = []
    filenames = []

    for filename in os.listdir(directory):
        if filename.endswith(ext):
            if ext == '.pdb':
                mol = Chem.MolFromPDBFile(os.path.join(directory, filename))
            elif ext == '.sdf':
                mol_supplier = Chem.SDMolSupplier(os.path.join(directory, filename))
                mol = next(mol_supplier)

            if mol is not None:
                # Split the filename by underscore, take the last part and remove the extension
                line_num = int(os.path.splitext(filename.rsplit('_', 1)[1])[0])

                library_line = linecache.getline(library_file, line_num).strip()
                split_line = library_line.split()
                library_line = f"{split_line[0]} https://zinc.docking.org/substances/ZINC{split_line[1]}"

                AllChem.Compute2DCoords(mol)

                molecules.append(mol)
                library_lines.append(library_line)
                filenames.append(filename)

    return molecules, library_lines, filenames

def create_pdf(molecules, library_lines, filenames, output_file):
    with PdfPages(output_file) as pdf:
        for i in range(0, len(molecules), 4):
            fig = plt.figure(figsize=(7.5, 10))

            for j in range(i, min(i+4, len(molecules))):
                mol = molecules[j]
                filename = filenames[j]
                library_line = library_lines[j]

                ax = fig.add_subplot(2, 2, j % 4 + 1)
                img = Draw.MolToImage(mol, size=(300, 300))
                ax.imshow(img)
                ax.axis('off')

                if (j - i) % 2 == 0:
                    ax.text(0.5, -0.1, f'{filename}\n{library_line}', fontsize=8, ha='center', va='top', transform=ax.transAxes)
                else:
                    ax.set_title(f'{filename}\n{library_line}', fontsize=8)
            pdf.savefig(bbox_inches='tight')
            plt.close()

def create_word(molecules, library_lines, filenames, output_file):
    document = Document()

    for i in range(0, len(molecules), 4):
        table = document.add_table(rows=2, cols=2)

        for j in range(i, min(i+4, len(molecules))):
            mol = molecules[j]
            filename = filenames[j]
            library_line = library_lines[j]

            row = j % 4 // 2
            col = j % 4 % 2
            cell = table.cell(row, col)

            # Add image to cell
            stream = BytesIO()
            img = Draw.MolToImage(mol, size=(600, 600))
            img.save(stream, format='png')
            stream.seek(0)
            cell.text = f'{filename}\n{library_line}'
            cell.add_paragraph().add_run().add_picture(stream, width=Inches(2.0))

    document.save(output_file)

def create_csv(library_lines, filenames, output_file):
    df = pd.DataFrame({
        'Filename': filenames,
        'Library_Line': library_lines
    })

    df.to_csv(output_file, index=False)

def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("directory", help="Path to the directory where the pdb/sdf and smi files are stored")
    parser.add_argument("output_prefix", help="Prefix for the output files")
    parser.add_argument("library_file", help="Path to the library file to read lines from")
    args = parser.parse_args()

    ext = '.pdb'
    pdf_output_file = args.output_prefix + '_output.pdf'
    docx_output_file = args.output_prefix + '_output.docx'
    csv_output_file = args.output_prefix + '_output.csv'
    molecules, library_lines, filenames = read_molecules(args.directory, ext, args.library_file)

    create_pdf(molecules, library_lines, filenames, pdf_output_file)
    create_word(molecules, library_lines, filenames, docx_output_file)
    create_csv(library_lines, filenames, csv_output_file)

if __name__ == "__main__":
    main()

